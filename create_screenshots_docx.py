import sys
import subprocess
import os

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('8. SCREENSHOTS', level=1)
    
    screenshots = [
        {
            "id": "8.1",
            "caption": 'In the above screen, the "AI Media Detector Login and OTP Authentication" is displayed.',
            "description": "Description: This screen displays the initial authentication gateway. Users can register with their email and verify their identity using a secure 6-digit OTP code before accessing the forensic dashboard."
        },
        {
            "id": "8.2",
            "caption": 'In the above screen, the main "AI Media Detector Dashboard" with media input options is displayed.',
            "description": "Description: This screen presents the primary interface. Users can choose to upload a direct media file (Video, Audio, Image) or paste public links to Instagram Reels and YouTube videos for forensic analysis."
        },
        {
            "id": "8.3",
            "caption": 'In the above screen, users can select their preferred "Detection Strategy" algorithm.',
            "description": 'Description: This screen allows users to choose between an "Auto Cascade" approach (utilizing Machine Learning + Heuristics) or a strict "Keyword Tag" metadata scanner to analyze the provided media.'
        },
        {
            "id": "8.4",
            "caption": 'In the above screen, the media uploading and loading analysis state is displayed.',
            "description": "Description: After providing a file or URL, the dashboard displays a loading state, indicating that the AI is actively scanning the specimen for deepfake artifacts, compression anomalies, and metadata inconsistencies."
        },
        {
            "id": "8.5",
            "caption": 'In the above screen, the "Analysis Complete" verdict for authentic media is displayed.',
            "description": 'Description: The dashboard displays a successful "Authentic Media Verified" verdict. It provides the probability scores, indicating a high confidence that the scanned media exhibits natural organic patterns without AI generation.'
        },
        {
            "id": "8.6",
            "caption": 'In the above screen, the "Analysis Complete" verdict for AI-generated media is displayed.',
            "description": 'Description: The dashboard displays an "AI Generated Content Detected" alert. The confidence metrics shift to indicate high AI probability, successfully flagging synthetic manipulation or altered metadata.'
        },
        {
            "id": "8.7",
            "caption": 'In the above screen, the "Advanced Forensic Analytics" benchmark panel is displayed.',
            "description": "Description: The dashboard presents comparative performance analytics, including Accuracy, Precision, Recall, and F1-Scores for both the Deep Learning (ViT) model and the Keyword Heuristics, alongside detailed Confusion Matrices."
        },
        {
            "id": "8.8",
            "caption": 'In the above screen, the "Forensic Heatmap (ELA/DCT Variance)" overlay is displayed.',
            "description": "Description: After completing the analysis, the user toggles the detailed forensic heatmap. This visual representation highlights localized manipulation hotspots in red and yellow, making it easy to observe forged regions."
        },
        {
            "id": "8.9",
            "caption": 'In the above screen, the ultimate "Multi-Modal Audio/Video Sync" detection flag is displayed.',
            "description": "Description: For deepfake MP4 videos, the dashboard demonstrates the advanced multi-modal evaluation, cross-referencing visual temporal jitter with audio spectral spikes to conclusively flag synthetic audio/video synchronization."
        }
    ]

    for item in screenshots:
        # Add a placeholder for the image
        p_img = doc.add_paragraph()
        p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_img = p_img.add_run('[ INSERT SCREENSHOT FOR {} HERE ]\n\n'.format(item['id']))
        run_img.font.size = Pt(14)
        run_img.font.bold = True
        
        # Add caption
        p_cap = doc.add_paragraph()
        run_cap = p_cap.add_run(f"{item['id']} {item['caption']}")
        run_cap.font.bold = True
        
        # Add description
        p_desc = doc.add_paragraph(item['description'] + '\n\n')
        
    doc.save(r'f:\Exp  AI Dect\AI_Detector_Screenshots.docx')

if __name__ == "__main__":
    create_document()
    print("Created AI_Detector_Screenshots.docx successfully.")
